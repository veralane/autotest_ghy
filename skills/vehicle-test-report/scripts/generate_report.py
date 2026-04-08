#!/usr/bin/env python3
"""
整车功能性能测试报告生成器

Usage:
    python generate_report.py --output report.docx [--data test_data.json] [--template template.docx]
"""

import argparse
import json
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请安装 python-docx: pip install python-docx")
    exit(1)


def set_cell_shading(cell, fill_color):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_with_number(doc, text, level=1):
    """添加带编号的标题"""
    heading = doc.add_heading(text, level=level)
    return heading


def create_report(data=None, template_path=None):
    """
    创建整车功能性能测试报告
    
    Args:
        data: 测试数据字典
        template_path: 模板文件路径
    
    Returns:
        Document对象
    """
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
    else:
        doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('整车功能性能测试报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 报告信息
    if data:
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"项目名称：{data.get('project_name', '【项目名称】')}\n")
        info_para.add_run(f"报告编号：{data.get('report_id', '【报告编号】')}\n")
        info_para.add_run(f"报告日期：{data.get('report_date', datetime.now().strftime('%Y-%m-%d'))}\n")
    
    doc.add_page_break()
    
    # 第一章 前言
    add_heading_with_number(doc, '第一章 前言', 1)
    
    if data and 'foreword' in data:
        doc.add_paragraph(data['foreword'])
    else:
        doc.add_paragraph('【前言内容】')
        doc.add_paragraph('本报告旨在记录和总结整车功能性能测试的执行情况及测试结果，为产品验收和质量评估提供依据。')
        doc.add_paragraph('本报告依据以下标准/规范编写：')
        doc.add_paragraph('• GB/T XXXXX-XXXX《XXXX》', style='List Bullet')
        doc.add_paragraph('• 企业内部测试规范', style='List Bullet')
    
    # 第二章 概述
    add_heading_with_number(doc, '第二章 概述', 1)
    
    if data and 'overview' in data:
        doc.add_paragraph(data['overview'])
    else:
        doc.add_paragraph('【概述内容】')
        
    add_heading_with_number(doc, '2.1 测试对象', 2)
    if data and 'test_object' in data:
        doc.add_paragraph(data['test_object'])
    else:
        doc.add_paragraph('【测试对象描述：车型、配置等信息】')
    
    add_heading_with_number(doc, '2.2 测试目标', 2)
    if data and 'test_objectives' in data:
        for obj in data['test_objectives']:
            doc.add_paragraph(f"• {obj}", style='List Bullet')
    else:
        doc.add_paragraph('【测试目标描述】')
    
    add_heading_with_number(doc, '2.3 测试周期', 2)
    if data and 'test_period' in data:
        doc.add_paragraph(data['test_period'])
    else:
        doc.add_paragraph('【测试起止时间】')
    
    # 第三章 测试版本说明
    add_heading_with_number(doc, '第三章 测试版本说明', 1)
    
    add_heading_with_number(doc, '3.1 测试版本信息', 2)
    
    # 版本信息表格
    version_table = doc.add_table(rows=5, cols=2)
    version_table.style = 'Table Grid'
    
    version_headers = ['项目', '版本信息']
    version_data = [
        ('软件版本', data.get('software_version', '【软件版本号】') if data else '【软件版本号】'),
        ('硬件版本', data.get('hardware_version', '【硬件版本号】') if data else '【硬件版本号】'),
        ('固件版本', data.get('firmware_version', '【固件版本号】') if data else '【固件版本号】'),
        ('版本变更说明', data.get('version_change', '【版本变更说明】') if data else '【版本变更说明】'),
    ]
    
    # 设置表头
    hdr_cells = version_table.rows[0].cells
    hdr_cells[0].text = version_headers[0]
    hdr_cells[1].text = version_headers[1]
    for cell in hdr_cells:
        set_cell_shading(cell, 'D9E2F3')
    
    # 填充数据
    for i, (item, value) in enumerate(version_data):
        row_cells = version_table.rows[i + 1].cells
        row_cells[0].text = item
        row_cells[1].text = value
    
    add_heading_with_number(doc, '3.2 测试环境描述', 2)
    
    # 测试环境表格
    env_table = doc.add_table(rows=5, cols=2)
    env_table.style = 'Table Grid'
    
    env_data = [
        ('测试场地', data.get('test_site', '【测试场地】') if data else '【测试场地】'),
        ('测试设备', data.get('test_equipment', '【测试设备】') if data else '【测试设备】'),
        ('测试工具', data.get('test_tools', '【测试工具】') if data else '【测试工具】'),
        ('环境条件', data.get('environment', '【温度、湿度等】') if data else '【温度、湿度等】'),
    ]
    
    hdr_cells = env_table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '描述'
    for cell in hdr_cells:
        set_cell_shading(cell, 'D9E2F3')
    
    for i, (item, value) in enumerate(env_data):
        row_cells = env_table.rows[i + 1].cells
        row_cells[0].text = item
        row_cells[1].text = value
    
    add_heading_with_number(doc, '3.3 引用的测试设计', 2)
    if data and 'test_design_refs' in data:
        for ref in data['test_design_refs']:
            doc.add_paragraph(f"• {ref}", style='List Bullet')
    else:
        doc.add_paragraph('【引用的测试设计文档】')
        doc.add_paragraph('• 测试用例文档：XXX-TEST-CASE-001', style='List Bullet')
        doc.add_paragraph('• 测试规范文档：XXX-TEST-SPEC-001', style='List Bullet')
    
    add_heading_with_number(doc, '3.4 测试通过标准', 2)
    if data and 'pass_criteria' in data:
        for criteria in data['pass_criteria']:
            doc.add_paragraph(f"• {criteria}", style='List Bullet')
    else:
        doc.add_paragraph('【测试通过标准】')
        doc.add_paragraph('• 功能通过准则：所有功能测试用例100%通过', style='List Bullet')
        doc.add_paragraph('• 性能指标要求：响应时间≤XXms', style='List Bullet')
    
    # 第四章 概要测试结论
    add_heading_with_number(doc, '第四章 概要测试结论', 1)
    
    add_heading_with_number(doc, '4.1 测试结论总结', 2)
    
    # 测试统计表格
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ('测试用例总数', str(data.get('total_cases', '【数量】')) if data else '【数量】'),
        ('通过用例数', str(data.get('passed_cases', '【数量】')) if data else '【数量】'),
        ('失败用例数', str(data.get('failed_cases', '【数量】')) if data else '【数量】'),
        ('测试结论', data.get('conclusion', '【通过/不通过/有条件通过】') if data else '【通过/不通过/有条件通过】'),
    ]
    
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = '统计项'
    hdr_cells[1].text = '数值'
    for cell in hdr_cells:
        set_cell_shading(cell, 'D9E2F3')
    
    for i, (item, value) in enumerate(summary_data):
        row_cells = summary_table.rows[i + 1].cells
        row_cells[0].text = item
        row_cells[1].text = value
    
    add_heading_with_number(doc, '4.2 关键风险和规避措施', 2)
    
    if data and 'risks' in data:
        risk_table = doc.add_table(rows=len(data['risks']) + 1, cols=4)
        risk_table.style = 'Table Grid'
        
        hdr_cells = risk_table.rows[0].cells
        headers = ['风险描述', '风险等级', '规避措施', '状态']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_shading(hdr_cells[i], 'D9E2F3')
        
        for i, risk in enumerate(data['risks']):
            row_cells = risk_table.rows[i + 1].cells
            row_cells[0].text = risk.get('description', '')
            row_cells[1].text = risk.get('level', '')
            row_cells[2].text = risk.get('mitigation', '')
            row_cells[3].text = risk.get('status', '')
    else:
        doc.add_paragraph('【关键风险和规避措施】')
        risk_table = doc.add_table(rows=2, cols=4)
        risk_table.style = 'Table Grid'
        
        hdr_cells = risk_table.rows[0].cells
        headers = ['风险描述', '风险等级', '规避措施', '状态']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_shading(hdr_cells[i], 'D9E2F3')
        
        for cell in risk_table.rows[1].cells:
            cell.text = '【待填写】'
    
    # 第五章 测试项目及结果
    add_heading_with_number(doc, '第五章 测试项目及结果', 1)
    
    add_heading_with_number(doc, '5.1 ABS测试结果', 2)
    
    # 5.1.1 干沥青直线制动
    add_heading_with_number(doc, '5.1.1 干沥青直线制动', 3)
    
    # 检查是否使用新格式数据
    if data and 'abs_straight_braking' in data:
        straight_data = data['abs_straight_braking']
        
        # 检查是否为新格式（包含test_items）
        if 'test_items' in straight_data:
            # 新格式：多车速、三次测试
            test_conditions = straight_data.get('test_conditions', {})
            
            # 测试条件
            doc.add_paragraph('测试条件：', style='Heading 4')
            cond_para = doc.add_paragraph()
            cond_para.add_run(f"• 测试路面：{test_conditions.get('测试路面', '【待填写】')}\n")
            cond_para.add_run(f"• 路面附着系数：{test_conditions.get('路面附着系数', '【待填写】')}\n")
            cond_para.add_run(f"• 测试环境：温度{test_conditions.get('测试温度', '【待填写】')}℃，湿度{test_conditions.get('测试湿度', '【待填写】')}%\n")
            
            doc.add_paragraph('测试结果记录表：', style='Heading 4')
            
            # 创建表格
            test_items = straight_data.get('test_items', [])
            # 计算总行数：每个车速4行（3次测试+平均）+ 1行要求值 + 1行表头
            total_rows = sum(len(item.get('test_runs', [])) + 1 for item in test_items) + 2
            
            table = doc.add_table(rows=total_rows, cols=9)
            table.style = 'Table Grid'
            
            # 表头
            headers = ['测试项目\n(车速)', '测试\n序号', '平均减速度\n(m/s²)', '制动距离\n(m)', 
                      '减速度相邻\n峰谷差值(m/s²)', '转向修正角\n(deg)', '车轮抱死\n时间(s)', 
                      '附着系数\n利用率(%)', '结论']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                set_cell_shading(hdr_cells[i], 'D9E2F3')
            
            # 填充数据
            row_idx = 1
            for item in test_items:
                speed = item.get('车速', '【待填写】')
                test_runs = item.get('test_runs', [])
                average = item.get('average', {})
                
                # 三次测试
                for run in test_runs:
                    cells = table.rows[row_idx].cells
                    cells[0].text = f"{speed} km/h" if run.get('序号') == 1 else ''
                    cells[1].text = str(run.get('序号', ''))
                    cells[2].text = str(run.get('平均减速度', ''))
                    cells[3].text = str(run.get('制动距离', ''))
                    cells[4].text = str(run.get('减速度峰谷差值', ''))
                    cells[5].text = str(run.get('转向修正角', ''))
                    cells[6].text = str(run.get('车轮抱死时间', ''))
                    cells[7].text = str(run.get('附着系数利用率', ''))
                    cells[8].text = run.get('结论', '')
                    row_idx += 1
                
                # 平均值
                cells = table.rows[row_idx].cells
                cells[0].text = ''
                cells[1].text = '平均'
                cells[2].text = str(average.get('平均减速度', ''))
                cells[3].text = str(average.get('制动距离', ''))
                cells[4].text = str(average.get('减速度峰谷差值', ''))
                cells[5].text = str(average.get('转向修正角', ''))
                cells[6].text = str(average.get('车轮抱死时间', ''))
                cells[7].text = str(average.get('附着系数利用率', ''))
                cells[8].text = average.get('结论', '')
                # 平均值行加粗
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                row_idx += 1
            
            # 要求值
            requirements = test_items[0].get('requirements', {}) if test_items else {}
            cells = table.rows[row_idx].cells
            cells[0].text = '要求值'
            cells[1].text = '-'
            cells[2].text = requirements.get('平均减速度', '【待填写】')
            cells[3].text = requirements.get('制动距离', '【待填写】')
            cells[4].text = requirements.get('减速度峰谷差值', '【待填写】')
            cells[5].text = requirements.get('转向修正角', '【待填写】')
            cells[6].text = requirements.get('车轮抱死时间', '【待填写】')
            cells[7].text = requirements.get('附着系数利用率', '【待填写】')
            cells[8].text = '通过'
            for cell in cells:
                set_cell_shading(cell, 'FFF2CC')
            
            # 主观评价
            subjective = straight_data.get('subjective_evaluation', '【待填写】')
            doc.add_paragraph(f"\n主观评价：{subjective}")
            
        else:
            # 旧格式兼容
            abs_test_fields = [
                '测试路面', '测试项目', '测试次数', '平均减速度(m/s²)',
                '制动距离(m)', '减速度相邻峰谷差值(m/s²)', '转向修正角(deg)',
                '单次循环车轮滑移率抱死时间(s)', '路面附着系数', '附着系数利用率',
                '主观评分', '主观评价', '结论'
            ]
            
            straight_table = doc.add_table(rows=len(abs_test_fields) + 1, cols=3)
            straight_table.style = 'Table Grid'
            
            hdr_cells = straight_table.rows[0].cells
            hdr_cells[0].text = '字段'
            hdr_cells[1].text = '测试要求'
            hdr_cells[2].text = '实测值'
            for cell in hdr_cells:
                set_cell_shading(cell, 'D9E2F3')
            
            for i, field in enumerate(abs_test_fields):
                row_cells = straight_table.rows[i + 1].cells
                row_cells[0].text = field
                row_cells[1].text = straight_data.get('requirements', {}).get(field, '【待填写】')
                row_cells[2].text = straight_data.get('measured', {}).get(field, '【待填写】')
    else:
        # 空白模板
        doc.add_paragraph('测试条件：', style='Heading 4')
        doc.add_paragraph('• 测试路面：【待填写】\n• 路面附着系数：【待填写】\n• 测试环境：温度【待填写】℃，湿度【待填写】%')
        
        doc.add_paragraph('测试结果记录表：', style='Heading 4')
        table = doc.add_table(rows=15, cols=9)
        table.style = 'Table Grid'
        
        headers = ['测试项目\n(车速)', '测试\n序号', '平均减速度\n(m/s²)', '制动距离\n(m)', 
                  '减速度相邻\n峰谷差值(m/s²)', '转向修正角\n(deg)', '车轮抱死\n时间(s)', 
                  '附着系数\n利用率(%)', '结论']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_shading(hdr_cells[i], 'D9E2F3')
        
        for row in table.rows[1:]:
            for cell in row.cells:
                cell.text = '【待填写】'
        
        doc.add_paragraph('\n主观评价：【待填写】')
    
    # 5.1.2 干沥青弯道制动
    add_heading_with_number(doc, '5.1.2 干沥青弯道制动', 3)
    
    if data and 'abs_curve_braking' in data:
        curve_data = data['abs_curve_braking']
        
        # 检查是否为新格式
        if 'test_items' in curve_data:
            test_conditions = curve_data.get('test_conditions', {})
            
            doc.add_paragraph('测试条件：', style='Heading 4')
            cond_para = doc.add_paragraph()
            cond_para.add_run(f"• 测试路面：{test_conditions.get('测试路面', '【待填写】')}\n")
            cond_para.add_run(f"• 弯道半径：{test_conditions.get('弯道半径', '【待填写】')}m\n")
            cond_para.add_run(f"• 路面附着系数：{test_conditions.get('路面附着系数', '【待填写】')}\n")
            cond_para.add_run(f"• 测试环境：温度{test_conditions.get('测试温度', '【待填写】')}℃，湿度{test_conditions.get('测试湿度', '【待填写】')}%\n")
            
            doc.add_paragraph('测试结果记录表：', style='Heading 4')
            
            test_items = curve_data.get('test_items', [])
            total_rows = sum(len(item.get('test_runs', [])) + 1 for item in test_items) + 2
            
            table = doc.add_table(rows=total_rows, cols=9)
            table.style = 'Table Grid'
            
            headers = ['测试项目\n(车速)', '测试\n序号', '平均减速度\n(m/s²)', '制动距离\n(m)', 
                      '减速度相邻\n峰谷差值(m/s²)', '转向修正角\n(deg)', '车轮抱死\n时间(s)', 
                      '附着系数\n利用率(%)', '结论']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                set_cell_shading(hdr_cells[i], 'D9E2F3')
            
            row_idx = 1
            for item in test_items:
                speed = item.get('车速', '【待填写】')
                test_runs = item.get('test_runs', [])
                average = item.get('average', {})
                
                for run in test_runs:
                    cells = table.rows[row_idx].cells
                    cells[0].text = f"{speed} km/h" if run.get('序号') == 1 else ''
                    cells[1].text = str(run.get('序号', ''))
                    cells[2].text = str(run.get('平均减速度', ''))
                    cells[3].text = str(run.get('制动距离', ''))
                    cells[4].text = str(run.get('减速度峰谷差值', ''))
                    cells[5].text = str(run.get('转向修正角', ''))
                    cells[6].text = str(run.get('车轮抱死时间', ''))
                    cells[7].text = str(run.get('附着系数利用率', ''))
                    cells[8].text = run.get('结论', '')
                    row_idx += 1
                
                cells = table.rows[row_idx].cells
                cells[0].text = ''
                cells[1].text = '平均'
                cells[2].text = str(average.get('平均减速度', ''))
                cells[3].text = str(average.get('制动距离', ''))
                cells[4].text = str(average.get('减速度峰谷差值', ''))
                cells[5].text = str(average.get('转向修正角', ''))
                cells[6].text = str(average.get('车轮抱死时间', ''))
                cells[7].text = str(average.get('附着系数利用率', ''))
                cells[8].text = average.get('结论', '')
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                row_idx += 1
            
            requirements = test_items[0].get('requirements', {}) if test_items else {}
            cells = table.rows[row_idx].cells
            cells[0].text = '要求值'
            cells[1].text = '-'
            cells[2].text = requirements.get('平均减速度', '【待填写】')
            cells[3].text = requirements.get('制动距离', '【待填写】')
            cells[4].text = requirements.get('减速度峰谷差值', '【待填写】')
            cells[5].text = requirements.get('转向修正角', '【待填写】')
            cells[6].text = requirements.get('车轮抱死时间', '【待填写】')
            cells[7].text = requirements.get('附着系数利用率', '【待填写】')
            cells[8].text = '通过'
            for cell in cells:
                set_cell_shading(cell, 'FFF2CC')
            
            subjective = curve_data.get('subjective_evaluation', '【待填写】')
            doc.add_paragraph(f"\n主观评价：{subjective}")
        else:
            # 旧格式兼容
            abs_test_fields = [
                '测试路面', '测试项目', '测试次数', '平均减速度(m/s²)',
                '制动距离(m)', '减速度相邻峰谷差值(m/s²)', '转向修正角(deg)',
                '单次循环车轮滑移率抱死时间(s)', '路面附着系数', '附着系数利用率',
                '主观评分', '主观评价', '结论'
            ]
            
            curve_table = doc.add_table(rows=len(abs_test_fields) + 1, cols=3)
            curve_table.style = 'Table Grid'
            
            hdr_cells = curve_table.rows[0].cells
            hdr_cells[0].text = '字段'
            hdr_cells[1].text = '测试要求'
            hdr_cells[2].text = '实测值'
            for cell in hdr_cells:
                set_cell_shading(cell, 'D9E2F3')
            
            for i, field in enumerate(abs_test_fields):
                row_cells = curve_table.rows[i + 1].cells
                row_cells[0].text = field
                row_cells[1].text = curve_data.get('requirements', {}).get(field, '【待填写】')
                row_cells[2].text = curve_data.get('measured', {}).get(field, '【待填写】')
    else:
        doc.add_paragraph('测试条件：', style='Heading 4')
        doc.add_paragraph('• 测试路面：【待填写】\n• 弯道半径：【待填写】m\n• 路面附着系数：【待填写】\n• 测试环境：温度【待填写】℃，湿度【待填写】%')
        
        doc.add_paragraph('测试结果记录表：', style='Heading 4')
        table = doc.add_table(rows=11, cols=9)
        table.style = 'Table Grid'
        
        headers = ['测试项目\n(车速)', '测试\n序号', '平均减速度\n(m/s²)', '制动距离\n(m)', 
                  '减速度相邻\n峰谷差值(m/s²)', '转向修正角\n(deg)', '车轮抱死\n时间(s)', 
                  '附着系数\n利用率(%)', '结论']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_shading(hdr_cells[i], 'D9E2F3')
        
        for row in table.rows[1:]:
            for cell in row.cells:
                cell.text = '【待填写】'
        
        doc.add_paragraph('\n主观评价：【待填写】')
    
    add_heading_with_number(doc, '5.2 TCS测试结果', 2)
    
    if data and 'tcs_results' in data:
        tcs_table = doc.add_table(rows=len(data['tcs_results']) + 1, cols=5)
        tcs_table.style = 'Table Grid'
        
        hdr_cells = tcs_table.rows[0].cells
        headers = ['用例编号', '测试项目', '预期结果', '实际结果', '结论']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_shading(hdr_cells[i], 'D9E2F3')
        
        for i, result in enumerate(data['tcs_results']):
            row_cells = tcs_table.rows[i + 1].cells
            row_cells[0].text = result.get('case_id', '')
            row_cells[1].text = result.get('test_item', '')
            row_cells[2].text = result.get('expected', '')
            row_cells[3].text = result.get('actual', '')
            row_cells[4].text = result.get('conclusion', '')
    else:
        doc.add_paragraph('【TCS测试结果】')
        tcs_table = doc.add_table(rows=4, cols=5)
        tcs_table.style = 'Table Grid'
        
        hdr_cells = tcs_table.rows[0].cells
        headers = ['用例编号', '测试项目', '预期结果', '实际结果', '结论']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_shading(hdr_cells[i], 'D9E2F3')
        
        for row in tcs_table.rows[1:]:
            for cell in row.cells:
                cell.text = '【待填写】'
    
    return doc


def main():
    parser = argparse.ArgumentParser(description='整车功能性能测试报告生成器')
    parser.add_argument('--output', '-o', default='vehicle_test_report.docx',
                        help='输出文件路径 (默认: vehicle_test_report.docx)')
    parser.add_argument('--data', '-d', help='测试数据JSON文件路径')
    parser.add_argument('--template', '-t', help='模板文件路径')
    
    args = parser.parse_args()
    
    # 加载测试数据
    data = None
    if args.data:
        with open(args.data, 'r', encoding='utf-8') as f:
            data = json.load(f)
    
    # 生成报告
    doc = create_report(data, args.template)
    
    # 保存文档
    doc.save(args.output)
    print(f'报告已生成: {args.output}')


if __name__ == '__main__':
    main()
